"""
文档转换工具 — FastAPI 后端
提供文件上传、转换任务管理、结果下载 API
"""

import os
import uuid
import secrets
import asyncio
import subprocess
import shutil
import logging

logger = logging.getLogger("doc-converter")
from pathlib import Path
from datetime import datetime
from typing import Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# PDF / Document libraries
from PyPDF2 import PdfReader, PdfWriter
from docx import Document

# ---- App Setup ----

app = FastAPI(title="文档转换工具", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---- Config ----

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# ---- In-memory task store ----

tasks: dict[str, dict] = {}

SUPPORTED_TYPES: dict[str, str] = {
    "pdf-to-word": "pdf",
    "pdf-to-excel": "pdf",
    "pdf-to-ppt": "pdf",
    "pdf-to-cad": "pdf",
    "cad-to-pdf": "dwg",
    "caj-to-pdf": "caj",
    "pdf-split": "pdf",
    "pdf-merge": "pdf",
    "pdf-to-searchable": "pdf",
    "pdf-scan-to-searchable": "pdf",
}

# ---- Models ----

class TaskStatus(BaseModel):
    task_id: str
    status: str  # pending, processing, done, error
    file_name: str
    percent: int = 0
    output_file: Optional[str] = None
    message: Optional[str] = None

# ---- Routes ----

@app.get("/api/health")
async def health():
    """健康检查"""
    return {"status": "ok", "version": "1.0.0"}


@app.post("/api/convert")
async def convert(
    files: list[UploadFile] = File(...),
    convert_type: str = Form(default="pdf-to-word"),
    output_format: str = Form(default="docx"),
    page_ranges: str = Form(default=""),
):
    """
    提交转换任务
    """
    if not files:
        raise HTTPException(status_code=400, detail="未选择文件")

    task_ids = []
    file_infos = []

    for file in files:
        if not file.filename:
            continue

        task_id = secrets.token_urlsafe(16)
        ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""

        # Save uploaded file
        safe_name = f"{task_id}_{file.filename}"
        upload_path = UPLOAD_DIR / safe_name
        content = await file.read()
        upload_path.write_bytes(content)

        file_infos.append({
            "task_id": task_id,
            "file_name": file.filename,
            "upload_path": str(upload_path),
        })

    # For merge: create one task with all file paths
    if convert_type == "pdf-merge":
        task_id = secrets.token_urlsafe(16)
        all_paths = [fi["upload_path"] for fi in file_infos]
        all_names = ", ".join(fi["file_name"] for fi in file_infos)
        tasks[task_id] = {
            "task_id": task_id,
            "status": "pending",
            "file_name": all_names if len(all_names) < 60 else f"{len(all_names)} 个文件",
            "percent": 0,
            "upload_path": all_paths[0] if all_paths else "",
            "upload_paths": all_paths,
            "convert_type": convert_type,
            "output_format": output_format,
            "page_ranges": page_ranges,
            "output_file": None,
            "message": "",
            "created_at": datetime.now().isoformat(),
        }
        task_ids.append(task_id)
    else:
        for fi in file_infos:
            tasks[fi["task_id"]] = {
                "task_id": fi["task_id"],
                "status": "pending",
                "file_name": fi["file_name"],
                "percent": 0,
                "upload_path": fi["upload_path"],
                "convert_type": convert_type,
                "output_format": output_format,
                "page_ranges": page_ranges,
                "output_file": None,
                "message": "",
                "created_at": datetime.now().isoformat(),
            }
            task_ids.append(fi["task_id"])

    # Start conversion in background
    for tid in task_ids:
        asyncio.create_task(run_conversion(tid))

    return {"task_ids": task_ids, "message": f"已接收 {len(task_ids)} 个任务"}


@app.get("/api/tasks/{task_id}")
async def get_task(task_id: str):
    """查询任务状态"""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return TaskStatus(**task)


@app.get("/api/download/{task_id}")
async def download(task_id: str):
    """下载转换结果"""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    if task["status"] != "done" or not task["output_file"]:
        raise HTTPException(status_code=400, detail="任务未完成或文件不存在")
    if not os.path.exists(task["output_file"]):
        raise HTTPException(status_code=404, detail="输出文件不存在")

    return FileResponse(
        path=task["output_file"],
        filename=os.path.basename(task["output_file"]),
        media_type="application/octet-stream",
    )


# ---- Conversion Engine ----

def find_libreoffice() -> Optional[str]:
    """查找 LibreOffice 可执行文件路径"""
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return shutil.which("soffice") or shutil.which("libreoffice")


def convert_via_libreoffice(input_path: str, output_dir: str, output_format: str) -> str:
    """使用 LibreOffice 转换（需要安装 LibreOffice）"""
    soffice = find_libreoffice()
    if not soffice:
        raise RuntimeError("LibreOffice 未安装")

    cmd = [
        soffice, "--headless",
        "--convert-to", output_format,
        "--outdir", output_dir,
        input_path,
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if proc.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {proc.stderr}")

    # LibreOffice names output as <stem>.<format>
    stem = Path(input_path).stem
    expected = os.path.join(output_dir, f"{stem}.{output_format}")
    if os.path.exists(expected):
        return expected
    # Try to find any newly created file
    candidates = list(Path(output_dir).glob(f"*.{output_format}"))
    if candidates:
        return str(max(candidates, key=lambda p: p.stat().st_mtime))
    raise RuntimeError("LibreOffice 输出文件未找到")


def convert_pdf_to_docx(input_path: str, output_path: str) -> None:
    """PDF → Word: 优先使用 pdf2docx 保留格式，fallback 到文本提取"""
    try:
        from pdf2docx import Converter
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
        return
    except ImportError:
        pass  # pdf2docx not installed, use fallback
    except Exception as e:
        logger.warning(f"pdf2docx failed, falling back to text extraction: {e}")

    # Fallback: plain text extraction via PyPDF2
    reader = PdfReader(input_path)
    doc = Document()

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            if i > 0:
                doc.add_page_break()
            doc.add_paragraph(text.strip())

    if not doc.paragraphs or all(not p.text.strip() for p in doc.paragraphs):
        doc.add_paragraph(f"[从 PDF 提取了 {len(reader.pages)} 页，但未检测到文本内容。")
        doc.add_paragraph("如果是扫描件，请使用 OCR 功能。]")

    doc.save(output_path)


def convert_pdf_split(input_path: str, output_dir: str, page_ranges: str = "") -> str:
    """PDF 拆分：按页码范围拆分为独立文件，打包为 zip"""
    import zipfile
    import re

    reader = PdfReader(input_path)
    total_pages = len(reader.pages)
    stem = Path(input_path).stem
    zip_path = os.path.join(output_dir, f"{stem}_split.zip")

    # Parse page ranges (1-indexed), e.g. "1-3, 5, 7-10"
    pages_to_extract: list[int] = []
    if page_ranges.strip():
        for part in re.split(r'[,，]\s*', page_ranges.strip()):
            part = part.strip()
            if not part:
                continue
            try:
                if '-' in part:
                    a, b = part.split('-', 1)
                    pages_to_extract.extend(range(int(a.strip()), int(b.strip()) + 1))
                else:
                    pages_to_extract.append(int(part))
            except ValueError:
                raise ValueError(f"无效的页码范围: '{part}'，请输入如 1-3, 5, 7-10 的格式")
        # Filter valid pages
        pages_to_extract = [p for p in pages_to_extract if 1 <= p <= total_pages]
        if not pages_to_extract:
            raise ValueError(f"页码范围 {page_ranges} 不在有效范围内（1-{total_pages}）")
    else:
        pages_to_extract = list(range(1, total_pages + 1))

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for page_num in pages_to_extract:
            writer = PdfWriter()
            writer.add_page(reader.pages[page_num - 1])
            page_path = os.path.join(output_dir, f"{stem}_p{page_num}.pdf")
            with open(page_path, 'wb') as f:
                writer.write(f)
            zf.write(page_path, f"page_{page_num}.pdf")
            os.remove(page_path)

    return zip_path


def convert_pdf_merge(input_paths: list[str], output_path: str) -> None:
    """PDF 合并"""
    merger = PdfWriter()
    for path in input_paths:
        reader = PdfReader(path)
        for page in reader.pages:
            merger.add_page(page)
    with open(output_path, 'wb') as f:
        merger.write(f)


async def run_conversion(task_id: str):
    """执行转换任务"""
    task = tasks[task_id]
    task["status"] = "processing"
    task["percent"] = 10

    try:
        input_path = task["upload_path"]
        convert_type = task["convert_type"]
        output_format = task["output_format"]
        page_ranges = task.get("page_ranges", "")
        output_name = Path(task["file_name"]).stem
        output_path = str(OUTPUT_DIR / f"{output_name}_{task_id}.{output_format}")

        task["percent"] = 20

        # ---- LibreOffice path (handles pdf-to-word/excel/ppt + cad/caj to pdf) ----
        lo_types = ("pdf-to-word", "pdf-to-excel", "pdf-to-ppt", "cad-to-pdf", "caj-to-pdf",
                    "pdf-to-cad", "pdf-to-searchable", "pdf-scan-to-searchable")
        if find_libreoffice() and convert_type in lo_types:
            task["message"] = "使用 LibreOffice 引擎..."
            task["percent"] = 30
            actual_output = await asyncio.to_thread(
                convert_via_libreoffice, input_path, str(OUTPUT_DIR), output_format
            )
            if actual_output != output_path:
                shutil.move(actual_output, output_path)

        # ---- Python engine path ----
        elif convert_type in ("pdf-to-word", "pdf-to-excel", "pdf-to-ppt"):
            task["message"] = "使用 Python 引擎提取文本..."
            task["percent"] = 40
            # Non-word types: force .docx since convert_pdf_to_docx outputs DOCX content
            if convert_type != "pdf-to-word":
                output_format = "docx"
                output_path = str(OUTPUT_DIR / f"{output_name}_{task_id}.docx")
                task["output_format"] = "docx"
                task["message"] = "已提取文本内容（安装 LibreOffice 可获得更好的格式保留效果）"
            await asyncio.to_thread(convert_pdf_to_docx, input_path, output_path)

        elif convert_type == "pdf-split":
            task["message"] = "拆分 PDF 页面..."
            task["percent"] = 40
            ext = "zip"
            output_path = str(OUTPUT_DIR / f"{output_name}_{task_id}.{ext}")
            zip_path = await asyncio.to_thread(convert_pdf_split, input_path, str(OUTPUT_DIR), page_ranges)
            output_path = zip_path

        elif convert_type == "pdf-merge":
            task["message"] = "合并 PDF..."
            task["percent"] = 40
            all_paths = task.get("upload_paths", [input_path])
            # Use stem from first file
            first_name = all_paths[0] if all_paths else input_path
            output_path = str(OUTPUT_DIR / f"merged_{task_id}.pdf")
            await asyncio.to_thread(convert_pdf_merge, all_paths, output_path)

        elif convert_type == "cad-to-pdf":
            task["message"] = "CAD 转 PDF 需要 LibreOffice 引擎"
            raise RuntimeError("需要安装 LibreOffice 才能转换 CAD 文件")

        elif convert_type == "caj-to-pdf":
            task["message"] = "CAJ 转 PDF 需要 LibreOffice 引擎"
            raise RuntimeError("需要安装 LibreOffice 才能转换 CAJ 文件")

        elif convert_type in ("pdf-to-searchable", "pdf-scan-to-searchable"):
            task["message"] = "OCR 需要安装 Tesseract"
            raise RuntimeError("需要安装 Tesseract OCR 引擎才能进行文字识别")

        else:
            task["message"] = "使用 Python 引擎..."
            task["percent"] = 40
            await asyncio.to_thread(convert_pdf_to_docx, input_path, output_path)

        task["percent"] = 90
        task["output_file"] = output_path
        task["status"] = "done"
        task["percent"] = 100
        task["message"] = "转换完成"

    except ValueError as e:
        task["status"] = "error"
        task["message"] = str(e)
        task["percent"] = 0
        logger.error(f"Task {task_id} validation error: {e}")
    except Exception as e:
        task["status"] = "error"
        task["message"] = "转换失败，请重试"
        task["percent"] = 0
        logger.error(f"Task {task_id} error: {e}", exc_info=True)


# ---- Startup ----

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8001, reload=True)
